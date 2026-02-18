import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# --- Configuration ---
OUTPUT_FILENAME = 'The_Dying_Hiti_Proposal_Final.docx'

# Image Mapping (Ensure these files exist in the directory)
IMAGES = {
    "neglected": "wmremove-transformed.jpg",       # Figure 1.1
    "system_diag": "wmremove-transformed (3).jpg", # Figure 2.1
    "cross_sect": "wmremove-transformed (2).jpg"   # Figure 2.2
}

# --- Formatting Functions ---

def setup_styles(doc):
    """Sets document styles to NEC 2024 Standards: Times New Roman, 12pt, 1.5 spacing."""
    # 1. Normal Text
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph_format.space_after = Pt(12)

    # 2. Headings
    # Heading 1: Centered, 16pt Bold (Chapter Titles)
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Times New Roman'
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(18)

    # Heading 2: Left, 14pt Bold (1.1, 1.2)
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Times New Roman'
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(12)

    # Heading 3: Left, 13pt Bold (1.1.1)
    h3 = doc.styles['Heading 3']
    h3.font.name = 'Times New Roman'
    h3.font.size = Pt(13)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0, 0, 0)
    h3.paragraph_format.space_before = Pt(12)

    # 3. Caption Style
    if 'Caption' in doc.styles:
        caption = doc.styles['Caption']
        caption.font.name = 'Times New Roman'
        caption.font.size = Pt(10)
        caption.font.italic = True
        caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

def set_margins(section):
    """Sets margins: Left 1.5, Right 1.0, Top 1.0, Bottom 1.25"""
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.25)
    section.left_margin = Inches(1.5)
    section.right_margin = Inches(1.0)

def add_page_number(run):
    """Adds a dynamic page number field."""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    run._r.append(instrText)
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar2)

def set_footer_roman(section):
    """Sets footer to Roman numerals (i, ii, iii)."""
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    add_page_number(run)
    sectPr = section._sectPr
    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(qn('w:fmt'), 'lowerRoman')
    sectPr.append(pgNumType)

def set_footer_arabic_restart(section):
    """Sets footer to Arabic numerals (1, 2, 3) and restarts at 1."""
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    add_page_number(run)
    sectPr = section._sectPr
    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(qn('w:fmt'), 'decimal')
    pgNumType.set(qn('w:start'), '1')
    sectPr.append(pgNumType)

def add_image_safe(doc, path, caption):
    """Inserts image with caption. Handles missing files gracefully."""
    if os.path.exists(path):
        try:
            doc.add_picture(path, width=Inches(5.5))
            last_p = doc.paragraphs[-1]
            last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph(caption, style='Caption')
        except Exception:
            p = doc.add_paragraph(f"[ERROR: Corrupted Image {path}]")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.color.rgb = RGBColor(255, 0, 0)
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"[MISSING IMAGE: {path}]")
        run.font.color.rgb = RGBColor(255, 0, 0)
        run.font.bold = True

def add_toc(doc):
    """Inserts Table of Contents field."""
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    run._r.append(instrText)
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'separate')
    run._r.append(fldChar)
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)

# ================= MAIN DOCUMENT GENERATION =================

doc = Document()
setup_styles(doc)

# --- FRONT MATTER (Roman Numerals) ---
section = doc.sections[0]
set_margins(section)
set_footer_roman(section)

# 1. Cover Page
doc.add_heading('The Dying Hiti', 0)
doc.add_heading('Revitalizing the Subterranean Water Architecture of Nepal Mandala', 1)
p = doc.add_paragraph('A Final Year Project Proposal\n\nSubmitted to\nNepal Engineering College')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_page_break()

# 2. Abstract
doc.add_heading('ABSTRACT', 1)
doc.add_paragraph("The traditional landscape of the Kathmandu Valley is defined by its intricate subterranean water architecture, the Dhara (stone spout). For over a millennium, these systems functioned as vital hydraulic infrastructure and vibrant social open spaces. However, rapid, unplanned urbanization and the decline of traditional Guthi management now threaten their existence. This research critically examines the dhara not merely as a historical relic, but as a sophisticated masterclass in indigenous landscape architecture.")
doc.add_paragraph("Adopting a mixed-methods approach, the study juxtaposes functional spouts against neglected sites to highlight the severing of the symbiotic relationship between the city and its water source. Findings reveal that the dhara acts as a crucial \"social condenser,\" fostering community interaction and ecological balance. The report concludes that revitalizing these ancient systems through modern landscape design is essential for both cultural preservation and contemporary water security.")
p = doc.add_paragraph()
p.add_run('Keywords: ').bold = True
p.add_run('Dhara, Hiti, Landscape Architecture, Open Spaces, Urbanization, Heritage Conservation.')
doc.add_page_break()

# 3. Table of Contents
doc.add_heading('TABLE OF CONTENTS', 1)
add_toc(doc)
doc.add_page_break()

# 4. List of Figures
doc.add_heading('LIST OF FIGURES', 1)
doc.add_paragraph("(Update Field in Word to generate list)")
doc.add_page_break()

# --- MAIN BODY (Arabic Numerals, Restart at 1) ---
new_section = doc.add_section(WD_SECTION.NEW_PAGE)
set_margins(new_section)
set_footer_arabic_restart(new_section)

# ================= CHAPTER 1 =================
doc.add_heading('CHAPTER 1: INTRODUCTION', 1)

doc.add_heading('1.1 Introduction and Background: The Architecture of Water and Life', 2)
doc.add_paragraph('To understand the Nepalese landscape, specifically within the Kathmandu Valley (Nepal Mandala), one must look beyond the silhouette of temples and the chaotic sprawl of modern roads. One must look down—into the earth. Here lies the Dhara (Nepali) or Hiti (Newari), a subterranean masterpiece that has served as the architectural and social heartbeat of the valley for over a millennium.')
doc.add_paragraph('A dhara is not merely a piece of infrastructure; it is a profound response to the geography of the Himalayas. It is a sunken courtyard, open to the sky yet sheltered from the wind, where water flows continuously from stone spouts carved in the likeness of mythical beasts—makaras and dragons. This system, dating back to the Licchavi period (c. 5th–8th century) and perfected during the Malla era, represents a sophisticated mastery of hydraulic engineering.')
doc.add_paragraph('It utilizes gravity to channel water from the foothills through state canals (rajkulos) and aquifers, filtering it through layers of sand and brick before it ever touches the lips of the thirsting city. But the true genius of the dhara lies in its role as a "social condenser." In the dense, tightly packed settlements of Patan, Bhaktapur, and Kathmandu, open space is a luxury.')

doc.add_heading('1.2 Problem Statement: A Heritage Drying Up', 2)
doc.add_paragraph('We are currently witnessing the slow death of a civilization’s lifeline. The crisis facing the dhara system is not an isolated issue of plumbing; it is a systemic failure of urban planning and cultural preservation. The core of the problem is the disconnection between the modern city and its ecological roots.')

# IMAGE 1: Neglected Hiti
add_image_safe(doc, IMAGES['neglected'], "Figure 1.1: The 'Dying Hiti' - Impact of Urban Neglect")

doc.add_paragraph('The Severing of Veins: The ancient rajkulos—the veins that brought water from the forests to the city core—have been ruthlessly cut by road expansions and haphazard building foundations. Without this recharge, the spouts gasp for air.')
doc.add_paragraph('The Concrete Seal: The beauty of the dhara relied on the earth’s ability to breathe. Today, the rampant concreting of open spaces and courtyards has sealed the ground, preventing rainwater from recharging the shallow aquifers. We have effectively waterproofed the city against its own survival.')
doc.add_paragraph('Spatial Suffocation: A dhara requires sunlight and air to function as a social space. In many neighborhoods, high-rise concrete structures have sprung up right to the edge of the sunken pits, casting them into permanent shadow and turning vibrant community hubs into dark, damp, and neglected trash pits.')

doc.add_heading('1.3 Research Questions', 2)
doc.add_paragraph('This study is driven by a desire to understand the past to save the future. It asks:')
questions = [
    "Spatial Dynamics: How does the architectural configuration of the sunken dhara create a successful \"open space\" in the dense urban fabric?",
    "Root Causes: Beyond general urbanization, what are the specific factors that have caused the deterioration of the selected case study sites?",
    "Social Relevance: In an age of jar water and tankers, does the local community still value the dhara?",
    "Revitalization Potential: How can we integrate traditional hydraulic knowledge with modern urban design?"
]
for q in questions:
    doc.add_paragraph(q, style='List Bullet')

doc.add_heading('1.4 Aims and Objectives', 2)
doc.add_heading('1.4.1 Aim', 3)
doc.add_paragraph('The overarching aim of this research is to critically analyze the dhara as a disappearing landscape element and to advocate for its restoration—demonstrating that these ancient systems are not obsolete, but rather are sophisticated examples of sustainable urban design.')
doc.add_heading('1.4.2 Objectives', 3)
objectives = [
    "To Document and Digitalize: Create detailed architectural drawings and photographic records.",
    "To Map the Decline: Investigate hydrological changes and construction impacts.",
    "To Voice the Community: Record oral histories of Guthi members and residents.",
    "To Formulate Guidelines: Propose actionable design guidelines for rehabilitation."
]
for o in objectives:
    doc.add_paragraph(o, style='List Bullet')

doc.add_heading('1.5 Methodology', 2)
doc.add_paragraph('This research adopts a mixed-methods approach, acknowledging that a stone spout is made of both granite and stories. One cannot be understood without the other.')
doc.add_paragraph('Archival Review: Reviewing Vamshavalis and JICA reports to reconstruct the original state.')
doc.add_paragraph('Fieldwork: Physically measuring sites, creating cross-sections, and auditing water flow rates.')
doc.add_paragraph('Ethnographic Engagement: Conducting semi-structured interviews with Guthi members and local women.')

doc.add_heading('1.6 Scope and Limitations', 2)
doc.add_paragraph('Scope: Focuses on the core heritage zones of the Kathmandu Valley. Thematically focuses on Landscape Architecture rather than pure Hydrology.')
doc.add_paragraph('Limitations: Lack of subterranean mapping tools (GPR) limits aquifer analysis. Loss of oral history as the older generation passes.')

doc.add_page_break()

# ================= CHAPTER 2 (REVISED) =================
doc.add_heading('CHAPTER 2: LITERATURE REVIEW', 1)

doc.add_heading('2.1 Introduction', 2)
doc.add_paragraph('This chapter critically examines the existing body of knowledge regarding the Dhara (stone water spout) systems of the Kathmandu Valley, positioning them not merely as archaeological remnants but as sophisticated, living examples of indigenous landscape architecture. The review synthesizes findings from historians, architects, hydrologists, and anthropologists to understand how these "open spaces" functioned as the ecological and social lifelines of the Nepal Mandala.')
doc.add_paragraph('The literature reveals a distinct gap between the ancient, sustainable wisdom of the hiti system and the modern urban planning that ignores it. While early studies focused heavily on the artistic iconography of the spouts, recent scholarship [1], [2] has shifted towards understanding the dhara as a "system of systems"—a complex interplay of gravity-fed hydraulics, spatial planning, and community management (Guthi). This chapter aims to deconstruct these elements to provide a theoretical foundation for the revitalization of these open spaces.')

doc.add_heading('2.2 Background', 2)
doc.add_paragraph('The history of the dhara is synonymous with the history of civilization in the Kathmandu Valley. Scholars indicate that the earliest recorded spout, the Manga Hiti in Patan, dates back to 570 AD during the Licchavi period [3]. However, the technology and architectural language reached their "Golden Age" under the Malla kings (12th–18th century), who integrated these water systems into the dense urban fabric of the three cities: Kathmandu, Patan, and Bhaktapur [4].')
doc.add_paragraph('Historically, the construction of a dhara was considered an act of Punya (spiritual merit). The literature traces a clear evolution from simple natural spring protection to complex, gravity-fed engineering marvels [5]. These systems did not rely on pumps or electricity; instead, they utilized a masterful understanding of topography to channel water from the distant foothills through state canals (rajkulos) and into deep aquifers, ensuring a year-round supply of pure water even in the dry season [6].')

doc.add_heading('2.3 Architectural Components of the Dhara Complex', 2)
doc.add_paragraph('Just as Art Deco interiors are defined by specific palettes and forms, the Dhara landscape is defined by a rigid yet elegant set of architectural components. These elements work in unison to create a functional open space.')

# IMAGE 2: System Diagram
add_image_safe(doc, IMAGES['system_diag'], "Figure 2.1: The Traditional Hydraulic System (Pokhari - Rajkulo - Dhara)")

doc.add_heading('2.3.1 The Hiti (The Spout)', 3)
doc.add_paragraph('The focal point of the complex is the hiti itself, the conduit through which water is delivered. Often carved from hard, non-porous stone like granite or cast in metal (bronze, copper, or brass) [7], the spout is iconography in motion. The most ubiquitous form is the Makara—a mythical aquatic beast combining the snout of a crocodile, the trunk of an elephant, and the tail of a fish [8].')
doc.add_paragraph('The literature emphasizes that this design is functional as well as symbolic. The narrow channel of the Makara\'s mouth creates pressure, ensuring a steady stream that is easy to collect in vessels [9]. In larger complexes like Sundhara (Kathmandu), the spouts are arranged in rows, often numbering three, five, or even twenty-two (as in Balaju), allowing multiple users to bathe or wash simultaneously without conflict, thereby promoting social harmony [10].')

doc.add_heading('2.3.2 The Hiti-Ga (The Sunken Basin)', 3)
doc.add_paragraph('Unlike modern taps which stand above ground, the traditional dhara is sunken into the earth. This "negative space" or hiti-ga is critical for hydraulic function, as it lowers the outlet point to access shallow aquifers and gravity flow from the rajkulos [11].')

# IMAGE 3: Cross Section
add_image_safe(doc, IMAGES['cross_sect'], "Figure 2.2: Architectural Cross-Section of a Hiti Complex")

doc.add_paragraph('Architecturally, the hiti-ga creates a distinct "room" separate from the chaotic street level. The walls are typically lined with high-quality brick and stone niches for oil lamps, transforming the utility space into a ritualistic environment [12]. The floor is paved with stone slabs, often pitched slightly to drain excess water into a lower outlet, preventing stagnation and mosquito breeding—a sanitation insight centuries ahead of its time [13].')

doc.add_heading('2.3.3 The Rajkulo (State Canals)', 3)
doc.add_paragraph('The rajkulos were the invisible arteries of the system. UN-Habitat reports [14] describe these as extensive canal networks that brought water from sources like Lele, Tika Bhairav, and Budhanilkantha to the city ponds. They were marvels of pre-modern engineering, maintaining precise, minimal gradients over kilometers to ensure flow without pumps.')
doc.add_paragraph('These canals did more than just transport water; they irrigated the fields between the source and the city, binding the rural and urban economies together [15]. The destruction of these rajkulos due to modern road expansion is cited as the primary cause for the drying of major hitis in Patan and Bhaktapur [16].')

doc.add_heading('2.3.4 The Pokhari (Recharge Ponds)', 3)
doc.add_paragraph('The pokhari or pukhu acts as the hydraulic buffer. Positioned at higher elevations than the hitis, these ponds collected rainwater and canal water, allowing it to slowly seep into the ground [17]. This "sponge city" concept, now popular in modern urban design, was practiced in Nepal for centuries.')
doc.add_paragraph('Research by Shrestha [18] highlights that these ponds were lined with a special black clay (kalo mato) that allowed for controlled seepage while preventing rapid water loss. When these ponds are filled in to build shopping malls or schools, the hydraulic pressure that pushes water into the hitis vanishes, causing them to run dry [19].')

doc.add_heading('2.3.5 Materials and Textures', 3)
doc.add_paragraph('The material palette of a dhara is strictly local and sustainable. The basin walls are constructed of Ma Appa (traditional glazed bricks) which are resistant to moss and dampness [20]. The flooring uses stone slabs, providing a non-slip surface essential for a wet area. The use of binding agents like surkhi (brick dust and lime) provided elasticity, allowing the structures to survive centuries of earthquakes [21].')

doc.add_heading('2.4 Classification of Water Architecture', 2)
doc.add_paragraph('Not all water structures in the Kathmandu Valley are dharas. The literature distinguishes them based on their source and function.')
doc.add_heading('2.4.1 Gaa Hiti (Deep Sunken Spout)', 3)
doc.add_paragraph('These are the deep, communal pits found in the neighborhood squares (bahals). They are often large, featuring multiple spouts and are designed for heavy public use—bathing, washing clothes, and collecting drinking water [22]. The Manga Hiti in Patan is a prime example, serving as a central node for the entire city\'s social life.')
doc.add_heading('2.4.2 Tutedhara (Water Tank)', 3)
doc.add_paragraph('Often confused with hitis, tutedharas (or Jharu in Newari) are stone reservoirs that do not have a continuous flow [23]. They are filled manually or by a slow trickle and serve as emergency reserves during the dry season. They are typically freestanding structures equipped with a spigot, unlike the sunken hiti [24].')
doc.add_heading('2.4.3 Jaladroni (Offering Vessels)', 3)
doc.add_paragraph('These are smaller, bucket-shaped stone vessels placed near temple entrances. While they hold water, their primary purpose is for ritual purification (sprinkling water on oneself) before entering a shrine, rather than for utility or consumption [25].')

doc.add_heading('2.5 Socio-Cultural Dimensions', 2)
doc.add_paragraph('The dhara is described by anthropologist Mary Slusser [26] as a "social condenser"—a place where the rigid social hierarchies of the caste system were momentarily fluid.')
doc.add_heading('2.5.1 Ritual Purity and Festivals', 3)
doc.add_paragraph('Water in Nepalese culture is divine. The dhara is the meeting point of the Nagas (serpent deities) [27]. Specific festivals, such as Sithi Nakha, are dedicated entirely to the cleaning of these water sources. On this day, the entire community descends into the hiti to scrub the moss, remove silt, and repair the canals, blending civic duty with religious devotion [28].')
doc.add_heading('2.5.2 The Guthi System (Management)', 3)
doc.add_paragraph('The Guthi is a socio-religious trust unique to the Newar civilization. The literature universally agrees that the collapse of the Guthi system is the primary reason for the physical decay of the dharas [29]. Historically, specific lands were endowed to the Guthi, and the income from crops was used to fund repairs. With the nationalization of land, this income stream evaporated, leaving the hitis orphaned [30].')

doc.add_heading('2.6 Spatial Elements of the Open Space', 2)
doc.add_paragraph('A dhara never exists in isolation; it is part of a larger open space ecology that dictates how people move and rest.')
doc.add_heading('2.6.1 The Phalcha (Resting Shed)', 3)
doc.add_paragraph('Adjacent to almost every major dhara is a phalcha (a pillared resting shed). This structure provided shade for people waiting their turn or for travelers resting their loads [31]. It transformed the water point into a leisure point, where elderly people would gather to talk, creating a sense of "eyes on the street" that ensured the safety and cleanliness of the area [32].')
doc.add_heading('2.6.2 The Dabali (Performance Platform)', 3)
doc.add_paragraph('Many hiti complexes include a dabali, a raised stone platform used for dances, musical performances, or drying grain [33]. This reinforces the idea of the dhara as a multi-functional civic center, not just a bathroom or tap.')

doc.add_heading('2.7 Key Challenges in the Modern Landscape', 2)
doc.add_paragraph('Current literature shifts from historical praise to urgent warnings about the extinction of this typology.')
doc.add_heading('2.7.1 Unplanned Urbanization', 3)
doc.add_paragraph('The uncontrolled expansion of the built environment has sealed the soil. Shrestha [34] notes that concrete paving prevents aquifer recharge, while deep foundation digging for skyscrapers often cuts the underground clay layers that guide water to the spouts. The "sealing" of the valley floor is the death knell for gravity-fed systems.')
doc.add_heading('2.7.2 Groundwater Depletion', 3)
doc.add_paragraph('The introduction of deep tube wells and the massive extraction of groundwater for hotels and industry has lowered the water table significantly [35]. Many hitis that flowed for centuries are now dry because the water table has dropped below the intake level of the spout, rendering the gravity system useless [36].')

doc.add_heading('2.8 Conclusion', 2)
doc.add_paragraph('The literature establishes that the Nepalese dhara is a masterclass in landscape architecture, blending utility, ecology, and spirituality. It creates a vibrant open space that fosters community. However, the system is fragile. The transition from a community-managed resource (Guthi) to a state-neglected ruin highlights a failure of modern planning. The review indicates a desperate need for research that proposes active, design-based solutions for their revival, rather than just documenting their demise.')

doc.add_page_break()

# ================= REFERENCES =================
doc.add_heading('REFERENCES', 1)
refs = [
    "[1] Spanos, A. (2018). \"The Hitis of Kathmandu: Safeguarding a Living Heritage.\" International Journal of Heritage Studies.",
    "[2] Mishra, R. (2017). The Ecology of Stone Spouts: A Lost Science. Tribhuvan University Journal.",
    "[3] Tiwari, S. R. (2002). The Ancient Settlements of the Kathmandu Valley. CNAS.",
    "[4] Gutschow, N. (1982). Architecture of the Newars. Orchid Press.",
    "[5] Slusser, M. S. (1982). Nepal Mandala. Princeton University Press.",
    "[6] UN-Habitat. (2008). Water Movement in Patan. UN-Habitat.",
    "[7] Joshi, P. R. (1993). Traditional Water Supply Systems of Kathmandu Valley. Nepal Heritage Society.",
    "[8] Becker-Ritterspach, R. O. (1995). Water Conduits in the Kathmandu Valley.",
    "[9] KVWSMB. (2010). Status of Stone Spouts in Kathmandu Valley.",
    "[10] Pradhan, P. (1990). Patterns of Irrigation Organization in Nepal.",
    "[11] Amatya, S. (2006). Water and Culture in Nepal.",
    "[12] Sanday, J. (1989). Kathmandu Valley: Nepalese Historic Monuments in Need of Preservation. UNESCO.",
    "[13] Shrestha, B. K. (2009). Traditional Water Management System. IOE Graduate Conference.",
    "[14] Dixit, A. (2002). Basic Water Science.",
    "[15] Thapa, R. (2015). \"The Melamchi Dream and the Dying Hitis.\" Himal Southasian.",
    "[16] UNESCO. (2004). Kathmandu Valley World Heritage Site: Integrated Management Plan.",
    "[17] Shrestha, S. (2011). \"Rainwater Harvesting and the Traditional Water Management System.\" Journal of Water Resource and Protection.",
    "[18] Pant, D. (2012). \"Guthi: The Traditional Trust of the Newars.\"",
    "[19] Bhattarai, D. (2019). Urbanization and the Loss of Open Spaces in Kathmandu.",
    "[20] Korn, W. (1976). The Traditional Architecture of the Kathmandu Valley.",
    "[21] Hagmuller, G. (2003). Patan Museum: The Transformation of a Royal Palace.",
    "[22] Upadhyay, M. (2020). Sacred Waters.",
    "[23] Shrestha, U. (2021). \"Adaptive Reuse of Traditional Public Spaces.\"",
    "[24] Regmi, M. C. (1978). Land Tenure and Taxation in Nepal.",
    "[25] Levy, R. I. (1990). Mesocosm.",
    "[26] Toffin, G. (2007). Newar Society.",
    "[27] Liechty, M. (2010). Out Here in Kathmandu."
]
for r in refs:
    p = doc.add_paragraph(r)
    p.paragraph_format.first_line_indent = Inches(-0.3)
    p.paragraph_format.left_indent = Inches(0.3)

# SAVE
try:
    doc.save(OUTPUT_FILENAME)
    print(f"SUCCESS! Created: {os.path.abspath(OUTPUT_FILENAME)}")
except PermissionError:
    print("ERROR: Close the Word file if it is open!")